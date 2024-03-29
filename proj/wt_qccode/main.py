# This is a sample Python script.

import io
import os.path
import fitz
from docx import Document
from docx.shared import Mm
from win32com.client import gencache
from win32com.client import constants, gencache

import qrcode


def createPdf(wordPath, pdfPath):
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(wordPath)
    doc.ExportAsFixedFormat(pdfPath,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)


def createQr(url):
    qrimg = qrcode.make(url)
    qrimg_bytes = io.BytesIO()
    qrimg.save(qrimg_bytes)
    return qrimg_bytes


def createQrtodocx(wordPath, charkey):
    # 显示XML文件
    # print(mydoc._element.xml)

    mydoc = Document(docx_filepath)

    for p in mydoc.paragraphs:
        for r in p.runs:
            # print(r.text)
            if charkey in r.text:
                isok = True
                qrimg = qrcode.make(url)
                qrimg_bytes = io.BytesIO()
                qrimg.save(qrimg_bytes)

                # r.clear()
                # r.add_text('补充内容')
                r.add_picture(qrimg_bytes, width=Mm(25))


def addImagetoPDF(x, y, px, img_bytes, pdf_filepath, filename, saveFilepach):
    # 定义坐标变量
    x, y = 50, 700
    px = 75
    rect = (x, y, x + px, y + px)  # 起始锚点横坐标，起始锚点纵坐标，结束锚点横坐标，结束锚点纵坐标
    img_xref = 0

    with fitz.open(pdf_filepath) as doc:
        # print(doc[0].getText())
        doc[0].insert_image(rect, stream=img)
        doc.save(saveFilepach+'/'+filename)


if __name__ == '__main__':
    # 设置基础变量
    filepath = 'C:/Users/kui/Desktop/mypdf'
    savefilepath = 'C:/Users/kui/Desktop/mypdf_ok'
    url = 'http://www.baidu.com'
    isok = False

    for filename in os.listdir(filepath):
        img = createQr(url)
        pdf_filepath = filepath + '/' + filename
        print(pdf_filepath)
        addImagetoPDF(50, 700, 75, img, pdf_filepath, filename, savefilepath)

    if isok:
        # mydoc.save(docx_filepath)

        if os.path.exists(pdf_filepath):
            os.remove(pdf_filepath)
        createPdf(docx_filepath, pdf_filepath)
        print('PDF文件已生成')
    else:
        print('没找到标签，请检查WORD文档是否有[二维码标记]信息')
